# import fitz  # PyMuPDF
# import docx

# def extract_text(file_path):
#     if file_path.endswith('.pdf'):
#         doc = fitz.open(file_path)
#         text = ""
#         for page in doc:
#             text += page.get_text()
#         return text
#     elif file_path.endswith('.docx'):
#         doc = docx.Document(file_path)
#         return "\n".join([p.text for p in doc.paragraphs])
#     else:
#         raise ValueError("Unsupported file type.")



import fitz  # PyMuPDF
import docx
import io

def extract_text(file_input):
    # If it's a file-like object (like from Streamlit)
    if hasattr(file_input, 'read'):
        if file_input.name.endswith('.pdf'):
            # Read PDF content into memory
            pdf_data = file_input.read()
            doc = fitz.open(stream=pdf_data, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        elif file_input.name.endswith('.docx'):
            doc = docx.Document(file_input)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError("Unsupported file type.")
    else:
        # If it's a string path (used in CLI)
        if file_input.endswith('.pdf'):
            doc = fitz.open(file_input)
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        elif file_input.endswith('.docx'):
            doc = docx.Document(file_input)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError("Unsupported file type.")
